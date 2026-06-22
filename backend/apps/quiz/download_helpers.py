import re
from io import BytesIO
from datetime import datetime
from zoneinfo import ZoneInfo
from html import unescape

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import json
import logging
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether,
    )
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

logger = logging.getLogger(__name__)


# ── Main flow (entry point first) ────────────────────────────────────────────

@csrf_exempt
@require_http_methods(["POST"])
def download_quiz_results(request):
    """Download quiz results as PDF, DOCX, or TXT."""
    try:
        data        = json.loads(request.body) if request.body else {}
        results     = data.get('results', {})
        file_format = data.get('format', 'pdf').lower()

        if not results:
            return JsonResponse({"error": "No results data provided"}, status=400)

        subject   = results.get('subject', 'Quiz')
        safe_name = _safe_filename(subject)
        tz        = ZoneInfo('Africa/Accra')
        ts        = datetime.now(tz).strftime('%Y%m%d_%H%M%S')
        base      = f"{safe_name}_Quiz_{ts}"

        if file_format == 'pdf':
            if not HAS_REPORTLAB:
                return JsonResponse(
                    {"error": "PDF generation unavailable. Please use TXT or DOCX."},
                    status=400,
                )
            buf = _build_pdf(results)
            response = HttpResponse(buf.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{base}.pdf"'
            return response

        elif file_format == 'docx':
            buf = _build_docx(results)
            response = HttpResponse(
                buf.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{base}.docx"'
            return response

        else:  # txt
            content  = _build_txt(results)
            response = HttpResponse(content, content_type='text/plain; charset=utf-8')
            response['Content-Disposition'] = f'attachment; filename="{base}.txt"'
            return response

    except Exception as e:
        logger.error(f"Error generating download: {e}", exc_info=True)
        return JsonResponse({"error": "Failed to generate file"}, status=500)


# ── Text helpers ──────────────────────────────────────────────────────────────

def _clean_html(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'<[^>]+>', '', text)
    text = unescape(text)
    return text.strip()


def _format_latex_readable(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\\[', '\n').replace('\\]', '\n').replace('$$', '\n')
    simple = {
        r'\\pi': 'π', r'\\alpha': 'α', r'\\beta': 'β', r'\\gamma': 'γ',
        r'\\delta': 'δ', r'\\theta': 'θ', r'\\lambda': 'λ', r'\\mu': 'μ',
        r'\\sigma': 'σ', r'\\omega': 'ω', r'\\Delta': 'Δ', r'\\Sigma': 'Σ',
        r'\\infty': '∞', r'\\pm': '±', r'\\times': '×', r'\\div': '÷',
        r'\\leq': '≤', r'\\geq': '≥', r'\\neq': '≠', r'\\approx': '≈',
        r'\\equiv': '≡', r'\\cdot': '·', r'\\rightarrow': '→',
        r'\\Rightarrow': '⇒', r'\\leftarrow': '←', r'\\Leftarrow': '⇐',
        r'\\leftrightarrow': '↔', r'\\partial': '∂', r'\\nabla': '∇',
        r'\\int': '∫', r'\\sum': '∑', r'\\prod': '∏', r'\\in': '∈',
        r'\\notin': '∉', r'\\subset': '⊂', r'\\subseteq': '⊆',
        r'\\cup': '∪', r'\\cap': '∩', r'\\emptyset': '∅',
        r'\\forall': '∀', r'\\exists': '∃', r'\\neg': '¬',
        r'\\wedge': '∧', r'\\vee': '∨',
    }
    for pat, rep in simple.items():
        text = text.replace(pat, rep)
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', text)
    text = re.sub(r'\^(\d)', lambda m: '⁰¹²³⁴⁵⁶⁷⁸⁹'[int(m.group(1))], text)
    text = re.sub(r'_(\d)',  lambda m: '₀₁₂₃₄₅₆₇₈₉'[int(m.group(1))],  text)
    text = text.replace('{', '').replace('}', '')
    text = re.sub(r'\\([a-zA-Z]+)', r'\1', text)
    return text.strip()


def _prep(text: str) -> str:
    return _format_latex_readable(_clean_html(text))


def _safe_filename(name: str, max_len: int = 180) -> str:
    if not name:
        return "Quiz_Results"
    s = re.sub(r'\s+', '_', str(name).strip())
    s = re.sub(r'[\\/:"*?<>|]+', '_', s)
    return s[:max_len] or "Quiz_Results"


# ── PDF ───────────────────────────────────────────────────────────────────────

def _build_pdf(results: dict) -> BytesIO:
    """Generate a professionally formatted PDF using ReportLab Platypus."""

    # Colour palette
    C = {
        'primary':       rl_colors.HexColor('#1d4ed8'),
        'primary_light': rl_colors.HexColor('#EEF2FF'),
        'answer':        rl_colors.HexColor('#059669'),
        'answer_bg':     rl_colors.HexColor('#ECFDF5'),
        'dark':          rl_colors.HexColor('#1E293B'),
        'body':          rl_colors.HexColor('#374151'),
        'gray':          rl_colors.HexColor('#64748B'),
        'light':         rl_colors.HexColor('#F8FAFC'),
        'border':        rl_colors.HexColor('#E2E8F0'),
        'expl_bg':       rl_colors.HexColor('#F1F5F9'),
        'white':         rl_colors.white,
        'indigo_text':   rl_colors.HexColor('#C7D2FE'),
    }

    # Layout constants
    PW     = letter[0]           # 612 pt
    MAR    = 0.7 * inch          # 50.4 pt
    CW     = PW - 2 * MAR        # ~511 pt content width
    CP     = 12                  # card padding (each side)
    INNER  = CW - 2 * CP         # inner card width

    BADGE_W   = 26
    Q_GAP     = 8
    Q_TEXT_W  = INNER - BADGE_W - Q_GAP

    OPT_INDENT = BADGE_W + Q_GAP   # 34 – aligns with question text
    OPT_INNER  = INNER - OPT_INDENT
    OPT_LBL_W  = 22
    OPT_TXT_W  = OPT_INNER - OPT_LBL_W

    ANS_LBL_W  = 52
    ANS_TXT_W  = INNER - ANS_LBL_W - 10

    EXPL_LBL_W = 70
    EXPL_TXT_W = INNER - EXPL_LBL_W - 6

    OPT_LABELS = ['A', 'B', 'C', 'D']

    buffer = BytesIO()

    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(C['gray'])
        canvas.drawString(MAR, 0.32 * inch, 'Generated by Ocasia')
        canvas.drawRightString(PW - MAR, 0.32 * inch, f'Page {doc.page}')
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        leftMargin=MAR, rightMargin=MAR,
        topMargin=MAR, bottomMargin=0.75 * inch,
    )

    def S(name, font='Helvetica', size=10, color=None, align=TA_LEFT,
          bold=False, italic=False, lm=1.4):
        fn = font
        if bold and italic:
            fn += '-BoldOblique'
        elif bold:
            fn += '-Bold'
        elif italic:
            fn += '-Oblique'
        return ParagraphStyle(name, fontName=fn, fontSize=size,
                              textColor=color or C['body'], alignment=align,
                              leading=size * lm)

    st = {
        'title':     S('title',     size=20, color=C['white'],       align=TA_CENTER, bold=True),
        'subtitle':  S('subtitle',  size=10, color=C['indigo_text'], align=TA_CENTER),
        'meta':      S('meta',      size=10, color=C['dark'],        align=TA_CENTER, lm=1.6),
        'q_num':     S('q_num',     size=11, color=C['white'],       align=TA_CENTER, bold=True),
        'q_text':    S('q_text',    size=11, color=C['dark'],        bold=True,  lm=1.45),
        'opt_label': S('opt_label', size=10, color=C['primary'],     align=TA_CENTER, bold=True),
        'opt_text':  S('opt_text',  size=10, color=C['body'],        lm=1.45),
        'ans_label': S('ans_label', size=8,  color=C['white'],       align=TA_CENTER, bold=True),
        'ans_text':  S('ans_text',  size=10, color=C['answer'],      bold=True, lm=1.45),
        'expl_head': S('expl_head', size=8,  color=C['gray'],        bold=True),
        'expl_text': S('expl_text', size=9,  color=C['gray'],        italic=True, lm=1.5),
        'footer':    S('footer',    size=9,  color=C['gray'],        align=TA_CENTER),
    }

    story = []

    subject    = results.get('subject', 'Quiz')
    difficulty = results.get('difficulty', 'Medium')
    src        = results.get('source_filename', '')
    details    = results.get('details', [])
    total      = len(details)
    tz         = ZoneInfo('Africa/Accra')
    date_str   = datetime.now(tz).strftime('%B %d, %Y')

    # ── Header banner ──────────────────────────────────────────────
    sub_parts = ['QUESTIONS & ANSWERS']
    if src:
        sub_parts.append(f'Source: {src}')

    banner = Table(
        [[Paragraph(subject.upper(), st['title'])],
         [Paragraph('  ·  '.join(sub_parts), st['subtitle'])]],
        colWidths=[CW],
    )
    banner.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, -1), C['primary']),
        ('LEFTPADDING',   (0, 0), (-1, -1), 16),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 16),
        ('TOPPADDING',    (0, 0), (0,  0),  18),
        ('BOTTOMPADDING', (0,-1), (-1, -1), 18),
        ('TOPPADDING',    (0, 1), (-1, -1),  4),
        ('BOTTOMPADDING', (0, 0), (-1, -2),  4),
    ]))
    story.append(banner)
    story.append(Spacer(1, 8))

    # ── Meta row ───────────────────────────────────────────────────
    meta = Table(
        [[Paragraph(f'<b>Difficulty</b><br/>{difficulty.title()}', st['meta']),
          Paragraph(f'<b>Questions</b><br/>{total}',               st['meta']),
          Paragraph(f'<b>Date</b><br/>{date_str}',                 st['meta'])]],
        colWidths=[CW / 3] * 3,
    )
    meta.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, -1), C['light']),
        ('GRID',          (0, 0), (-1, -1), 0.5, C['border']),
        ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING',    (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(meta)
    story.append(Spacer(1, 22))

    # ── Questions ──────────────────────────────────────────────────
    for idx, detail in enumerate(details, 1):
        q_text  = _prep(detail.get('question', ''))
        options = detail.get('options', [])
        ans     = _prep(detail.get('correct_answer', 'N/A'))
        expl    = _prep(detail.get('explanation', ''))

        rows       = []
        row_styles = []

        # ─ Question header: badge + text ──────────────────────────
        badge = Table([[Paragraph(str(idx), st['q_num'])]], colWidths=[BADGE_W])
        badge.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), C['primary']),
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING',   (0, 0), (-1, -1), 2),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 2),
        ]))
        q_hdr = Table(
            [[badge, Paragraph(q_text, st['q_text'])]],
            colWidths=[BADGE_W, Q_TEXT_W],
        )
        q_hdr.setStyle(TableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING',   (0, 0), (-1, -1), 0),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('LEFTPADDING',   (1, 0), (1,  0),  Q_GAP),
        ]))
        q_row_i = len(rows)
        rows.append([q_hdr])
        row_styles += [
            ('TOPPADDING',    (0, q_row_i), (0, q_row_i), 14),
            ('BOTTOMPADDING', (0, q_row_i), (0, q_row_i),  8),
        ]

        # ─ Options (MCQ only) ─────────────────────────────────────
        if options:
            opt_rows = []
            for i, opt in enumerate(options[:4]):
                lbl = Table([[Paragraph(OPT_LABELS[i], st['opt_label'])]], colWidths=[OPT_LBL_W])
                lbl.setStyle(TableStyle([
                    ('BACKGROUND',    (0, 0), (-1, -1), C['primary_light']),
                    ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING',    (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('LEFTPADDING',   (0, 0), (-1, -1), 2),
                    ('RIGHTPADDING',  (0, 0), (-1, -1), 2),
                ]))
                opt_rows.append([lbl, Paragraph(_prep(str(opt)), st['opt_text'])])

            opts_tbl = Table(opt_rows, colWidths=[OPT_LBL_W, OPT_TXT_W])
            opts_tbl.setStyle(TableStyle([
                ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING',   (0, 0), (0,  -1),  0),
                ('RIGHTPADDING',  (0, 0), (0,  -1),  0),
                ('LEFTPADDING',   (1, 0), (1,  -1),  6),
                ('RIGHTPADDING',  (1, 0), (1,  -1),  0),
                ('TOPPADDING',    (0, 0), (-1, -1),  4),
                ('BOTTOMPADDING', (0, 0), (-1, -1),  4),
                ('LINEBELOW',     (0, 0), (-1, -2),  0.25, C['border']),
            ]))

            # Indent by wrapping in 2-col outer: [empty | opts]
            opts_outer = Table([['', opts_tbl]], colWidths=[OPT_INDENT, OPT_INNER])
            opts_outer.setStyle(TableStyle([
                ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                ('TOPPADDING',    (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ]))

            opt_row_i = len(rows)
            rows.append([opts_outer])
            row_styles += [
                ('TOPPADDING',    (0, opt_row_i), (0, opt_row_i), 2),
                ('BOTTOMPADDING', (0, opt_row_i), (0, opt_row_i), 6),
            ]

        # ─ Answer ─────────────────────────────────────────────────
        ans_lbl = Table([[Paragraph('ANSWER', st['ans_label'])]], colWidths=[ANS_LBL_W])
        ans_lbl.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), C['answer']),
            ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING',   (0, 0), (-1, -1), 4),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 4),
        ]))
        ans_tbl = Table(
            [[ans_lbl, Paragraph(ans, st['ans_text'])]],
            colWidths=[ANS_LBL_W, ANS_TXT_W],
        )
        ans_tbl.setStyle(TableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING',   (0, 0), (-1, -1), 0),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ('LEFTPADDING',   (1, 0), (1,  0),  10),
        ]))

        ans_row_i = len(rows)
        rows.append([ans_tbl])
        row_styles += [
            ('BACKGROUND',    (0, ans_row_i), (0, ans_row_i), C['answer_bg']),
            ('TOPPADDING',    (0, ans_row_i), (0, ans_row_i),  8),
            ('BOTTOMPADDING', (0, ans_row_i), (0, ans_row_i),  8),
        ]

        # ─ Explanation ────────────────────────────────────────────
        if expl:
            expl_tbl = Table(
                [[Paragraph('Explanation:', st['expl_head']),
                  Paragraph(expl, st['expl_text'])]],
                colWidths=[EXPL_LBL_W, EXPL_TXT_W],
            )
            expl_tbl.setStyle(TableStyle([
                ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                ('RIGHTPADDING',  (0, 0), (0,  -1),  6),
                ('LEFTPADDING',   (1, 0), (1,  -1),  0),
                ('TOPPADDING',    (0, 0), (-1, -1),  0),
                ('BOTTOMPADDING', (0, 0), (-1, -1),  0),
            ]))

            expl_row_i = len(rows)
            rows.append([expl_tbl])
            row_styles += [
                ('BACKGROUND',    (0, expl_row_i), (0, expl_row_i), C['expl_bg']),
                ('TOPPADDING',    (0, expl_row_i), (0, expl_row_i),  8),
                ('BOTTOMPADDING', (0, expl_row_i), (0, expl_row_i),  8),
            ]

        # ─ Card wrapper ───────────────────────────────────────────
        last_i = len(rows) - 1
        base_styles = [
            ('BOX',           (0, 0),     (-1, -1), 0.75, C['border']),
            ('BACKGROUND',    (0, 0),     (-1, -1), C['white']),
            ('LEFTPADDING',   (0, 0),     (-1, -1), CP),
            ('RIGHTPADDING',  (0, 0),     (-1, -1), CP),
            ('TOPPADDING',    (0, 0),     (-1, -1),  4),
            ('BOTTOMPADDING', (0, 0),     (-1, -1),  4),
            ('BOTTOMPADDING', (0, last_i),(0, last_i), 14),
        ]
        card = Table(rows, colWidths=[CW])
        card.setStyle(TableStyle(base_styles + row_styles))
        story.append(KeepTogether([card, Spacer(1, 12)]))

    # ── Footer ─────────────────────────────────────────────────────
    story.append(HRFlowable(width='100%', thickness=0.5, color=C['border']))
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        'Generated by <b>Ocasia</b>  ·  https://ocasia.live',
        st['footer'],
    ))

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    buffer.seek(0)
    return buffer


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _shade_para(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)


def _add_hr(doc, color_hex='E2E8F0'):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    return para


def _set_table_full_width(table):
    tbl  = table._tbl
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '5000')
    tblW.set(qn('w:type'), 'pct')
    table._tbl.tblPr.append(tblW)


def _remove_cell_borders(cell):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBdr.append(b)
    tcPr.append(tcBdr)


def _build_docx(results: dict) -> BytesIO:
    """Generate a professionally formatted DOCX."""
    OPT_LABELS = ['A', 'B', 'C', 'D']

    RGB = {
        'primary':    RGBColor(0x4F, 0x46, 0xE5),
        'answer':     RGBColor(0x05, 0x96, 0x69),
        'dark':       RGBColor(0x1E, 0x29, 0x3B),
        'body':       RGBColor(0x37, 0x41, 0x51),
        'gray':       RGBColor(0x64, 0x74, 0x8B),
        'white':      RGBColor(0xFF, 0xFF, 0xFF),
        'indigo_sub': RGBColor(0xC7, 0xD2, 0xFE),
    }
    HEX = {
        'primary':    '4F46E5',
        'primary_lt': 'EEF2FF',
        'answer':     '059669',
        'answer_bg':  'ECFDF5',
        'expl_bg':    'F1F5F9',
        'light':      'F8FAFC',
        'border':     'E2E8F0',
        'white':      'FFFFFF',
    }

    subject    = results.get('subject', 'Quiz')
    difficulty = results.get('difficulty', 'Medium')
    src        = results.get('source_filename', '')
    details    = results.get('details', [])
    total      = len(details)
    tz         = ZoneInfo('Africa/Accra')
    date_str   = datetime.now(tz).strftime('%B %d, %Y')

    doc = docx.Document()

    for section in doc.sections:
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(2)

    # ── Header banner ──────────────────────────────────────────────
    banner_tbl = doc.add_table(rows=2, cols=1)
    _set_table_full_width(banner_tbl)
    banner_tbl.style = 'Table Grid'

    title_cell = banner_tbl.rows[0].cells[0]
    _set_cell_bg(title_cell, HEX['primary'])
    _remove_cell_borders(title_cell)
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.paragraph_format.space_before = Pt(14)
    title_para.paragraph_format.space_after  = Pt(2)
    tr = title_para.add_run(subject.upper())
    tr.bold = True
    tr.font.size      = Pt(20)
    tr.font.color.rgb = RGB['white']

    sub_cell = banner_tbl.rows[1].cells[0]
    _set_cell_bg(sub_cell, HEX['primary'])
    _remove_cell_borders(sub_cell)
    sub_para = sub_cell.paragraphs[0]
    sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after  = Pt(14)
    sub_parts = ['QUESTIONS & ANSWERS']
    if src:
        sub_parts.append(f'Source: {src}')
    sr = sub_para.add_run('  ·  '.join(sub_parts))
    sr.font.size      = Pt(10)
    sr.font.color.rgb = RGB['indigo_sub']

    doc.add_paragraph()

    # ── Meta row ───────────────────────────────────────────────────
    meta_tbl = doc.add_table(rows=1, cols=3)
    _set_table_full_width(meta_tbl)
    meta_tbl.style = 'Table Grid'
    for col_i, (label, value) in enumerate([
        ('Difficulty', difficulty.title()),
        ('Questions',  str(total)),
        ('Date',       date_str),
    ]):
        cell = meta_tbl.rows[0].cells[col_i]
        _set_cell_bg(cell, HEX['light'])
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        lr = p.add_run(f'{label}\n')
        lr.bold = True
        lr.font.size      = Pt(8)
        lr.font.color.rgb = RGB['gray']
        vr = p.add_run(value)
        vr.bold = True
        vr.font.size      = Pt(11)
        vr.font.color.rgb = RGB['dark']

    doc.add_paragraph()

    # ── Questions ──────────────────────────────────────────────────
    for idx, detail in enumerate(details, 1):
        q_text  = _prep(detail.get('question', ''))
        options = detail.get('options', [])
        ans     = _prep(detail.get('correct_answer', 'N/A'))
        expl    = _prep(detail.get('explanation', ''))

        # ─ Question text ──────────────────────────────────────────
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(4)
        q_para.paragraph_format.space_after  = Pt(4)
        nr = q_para.add_run(f'Q{idx}.  ')
        nr.bold = True
        nr.font.size      = Pt(11)
        nr.font.color.rgb = RGB['primary']
        tr2 = q_para.add_run(q_text)
        tr2.bold = True
        tr2.font.size      = Pt(11)
        tr2.font.color.rgb = RGB['dark']

        # ─ Options ────────────────────────────────────────────────
        if options:
            opt_tbl = doc.add_table(rows=len(options[:4]), cols=2)
            _set_table_full_width(opt_tbl)
            opt_tbl.style = 'Table Grid'
            for i, opt in enumerate(options[:4]):
                row = opt_tbl.rows[i]

                lbl_cell = row.cells[0]
                lbl_cell.width = Inches(0.4)
                _set_cell_bg(lbl_cell, HEX['primary_lt'])
                lp = lbl_cell.paragraphs[0]
                lp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                lp.paragraph_format.space_before = Pt(4)
                lp.paragraph_format.space_after  = Pt(4)
                llr = lp.add_run(OPT_LABELS[i])
                llr.bold = True
                llr.font.size      = Pt(10)
                llr.font.color.rgb = RGB['primary']

                txt_cell = row.cells[1]
                _set_cell_bg(txt_cell, HEX['white'])
                tp = txt_cell.paragraphs[0]
                tp.paragraph_format.space_before = Pt(4)
                tp.paragraph_format.space_after  = Pt(4)
                tp.paragraph_format.left_indent  = Inches(0.05)
                otr = tp.add_run(_prep(str(opt)))
                otr.font.size      = Pt(10)
                otr.font.color.rgb = RGB['body']

            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)

        # ─ Answer ─────────────────────────────────────────────────
        ans_para = doc.add_paragraph()
        _shade_para(ans_para, HEX['answer_bg'])
        ans_para.paragraph_format.space_before = Pt(2)
        ans_para.paragraph_format.space_after  = Pt(2)
        ans_para.paragraph_format.left_indent  = Inches(0.05)
        al = ans_para.add_run('ANSWER:  ')
        al.bold = True
        al.font.size      = Pt(9)
        al.font.color.rgb = RGB['answer']
        av = ans_para.add_run(ans)
        av.bold = True
        av.font.size      = Pt(10)
        av.font.color.rgb = RGB['answer']

        # ─ Explanation ────────────────────────────────────────────
        if expl:
            expl_para = doc.add_paragraph()
            _shade_para(expl_para, HEX['expl_bg'])
            expl_para.paragraph_format.space_before = Pt(1)
            expl_para.paragraph_format.space_after  = Pt(4)
            expl_para.paragraph_format.left_indent  = Inches(0.05)
            el = expl_para.add_run('Explanation:  ')
            el.bold = True
            el.font.size      = Pt(9)
            el.font.color.rgb = RGB['gray']
            ev = expl_para.add_run(expl)
            ev.italic = True
            ev.font.size      = Pt(9)
            ev.font.color.rgb = RGB['gray']

        # ─ Divider (not after last) ───────────────────────────────
        if idx < total:
            _add_hr(doc, HEX['border'])
            sp2 = doc.add_paragraph()
            sp2.paragraph_format.space_after = Pt(2)

    # ── Footer ─────────────────────────────────────────────────────
    doc.add_paragraph()
    _add_hr(doc, HEX['border'])
    foot = doc.add_paragraph()
    foot.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    foot.paragraph_format.space_before = Pt(6)
    fr = foot.add_run('Generated by Ocasia  ·  https://ocasia.live')
    fr.font.size      = Pt(9)
    fr.font.color.rgb = RGB['gray']

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── TXT ───────────────────────────────────────────────────────────────────────

def _build_txt(results: dict) -> str:
    details    = results.get('details', [])
    subject    = results.get('subject', 'Quiz')
    difficulty = results.get('difficulty', 'Random')
    src        = results.get('source_filename', '')
    tz         = ZoneInfo('Africa/Accra')
    ts         = datetime.now(tz).strftime('%Y-%m-%d %H:%M:%S %Z')
    OPT_LABELS = ['A', 'B', 'C', 'D']

    lines = [
        '=' * 72,
        f'  {subject.upper()} — QUIZ QUESTIONS & ANSWERS',
        '=' * 72,
        f'  Difficulty : {difficulty.title()}',
        f'  Questions  : {len(details)}',
        f'  Generated  : {ts}',
    ]
    if src:
        lines.append(f'  Source     : {src}')
    lines += ['=' * 72, '']

    for idx, detail in enumerate(details, 1):
        q_text  = _prep(detail.get('question', ''))
        options = detail.get('options', [])
        ans     = _prep(detail.get('correct_answer', 'N/A'))
        expl    = _prep(detail.get('explanation', ''))

        lines.append(f'Question {idx}')
        lines.append(f'  {q_text}')
        lines.append('')

        if options:
            for i, opt in enumerate(options[:4]):
                lines.append(f'  {OPT_LABELS[i]}.  {_prep(str(opt))}')
            lines.append('')

        lines.append(f'  Answer: {ans}')
        if expl:
            lines.append(f'  Explanation: {expl}')
        lines += ['', '-' * 72, '']

    lines += ['', 'Generated by Ocasia  —  https://ocasia.live']
    return '\n'.join(lines)


